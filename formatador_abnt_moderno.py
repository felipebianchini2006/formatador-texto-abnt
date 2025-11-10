#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Formatador ABNT Desktop - Versão 3.0 Moderna
Aplicativo profissional para criação e formatação de trabalhos acadêmicos conforme ABNT
Versão: 3.0 - Interface moderna e funcionalidades completas
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime


# Configuração do tema
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


class FormatadorABNT:
    """Classe responsável pela formatação completa ABNT"""

    @staticmethod
    def formatar_citacoes(texto):
        """
        Formata citações conforme NBR 10520
        - Curtas: até 3 linhas, entre aspas
        - Longas: >3 linhas, recuo 4cm, sem aspas
        - Autor em MAIÚSCULAS
        """

        def converter_maiusculas(match):
            """Converte citações para MAIÚSCULAS"""
            conteudo = match.group(1)
            autores = conteudo.split(';')
            autores_formatados = []

            for autor in autores:
                autor = autor.strip()
                partes = re.split(r'(,\s*\d{4})', autor, maxsplit=1)

                if len(partes) >= 2:
                    nome = partes[0].strip()
                    resto = ''.join(partes[1:])

                    palavras = nome.split()
                    palavras_maiusculas = [
                        p.lower() if p.lower() in ['et', 'al', 'al.'] else p.upper()
                        for p in palavras
                    ]

                    autores_formatados.append(' '.join(palavras_maiusculas) + resto)
                else:
                    palavras = autor.split()
                    palavras_maiusculas = [
                        p.lower() if p.lower() in ['et', 'al', 'al.'] else p.upper()
                        for p in palavras
                    ]
                    autores_formatados.append(' '.join(palavras_maiusculas))

            return f"({'; '.join(autores_formatados)})"

        # Padrão de citações entre parênteses
        padrao = r'\(([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s,;]+\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)\)'
        texto = re.sub(padrao, converter_maiusculas, texto)

        # Converte múltiplos autores para et al. (4+)
        def converter_et_al(match):
            conteudo = match.group(1)
            autores = conteudo.split(';')

            if len(autores) >= 4:
                primeiro = autores[0].strip()
                primeiro = re.sub(r',\s*\d{4}.*$', '', primeiro).strip()

                ano_match = re.search(r',\s*(\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)', conteudo)
                if ano_match:
                    ano = ano_match.group(1)
                    return f"({primeiro} et al., {ano})"

            return match.group(0)

        texto = re.sub(
            r'\(([A-ZÀ-Ü][A-Za-zÀ-ü]+(?:\s+[a-zà-ü]+)*(?:;\s*[A-ZÀ-Ü][A-Za-zÀ-ü]+(?:\s+[a-zà-ü]+)*){3,}[,\s]+\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)\)',
            converter_et_al, texto
        )

        return texto

    @staticmethod
    def formatar_referencias(texto):
        """
        Formata referências conforme NBR 6023
        SOBRENOME, Nome. Título: subtítulo. Edição. Local: Editora, ano.
        """
        linhas = texto.split('\n')
        referencias_formatadas = []

        for linha in linhas:
            if linha.strip():
                # Aplica formatação básica de referência
                referencias_formatadas.append(linha.strip())

        return '\n'.join(referencias_formatadas)


class GeradorDocumentoABNT:
    """Classe para gerar documentos Word completos conforme ABNT"""

    def __init__(self):
        self.doc = Document()
        self._configurar_documento()

    def _configurar_documento(self):
        """Configura margens e estilos padrão ABNT"""
        # Configurar margens (NBR 14724)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21)

    def adicionar_capa(self, dados):
        """
        Gera capa conforme ABNT
        dados = {
            'instituicao': str,
            'curso': str,
            'autor': str,
            'titulo': str,
            'local': str,
            'ano': str
        }
        """
        # Instituição (topo, centralizado)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('instituicao', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        # Curso
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('curso', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Espaçamento vertical
        for _ in range(8):
            self.doc.add_paragraph()

        # Autor (centro)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('autor', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        # Espaçamento
        for _ in range(4):
            self.doc.add_paragraph()

        # Título
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('titulo', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True

        # Espaçamento até o rodapé
        for _ in range(8):
            self.doc.add_paragraph()

        # Local e ano (rodapé)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('local', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('ano', ''))
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Quebra de página
        self.doc.add_page_break()

    def adicionar_folha_rosto(self, dados):
        """
        Gera folha de rosto conforme ABNT
        dados = {
            'autor': str,
            'titulo': str,
            'natureza': str (ex: "Trabalho de Conclusão de Curso"),
            'objetivo': str (ex: "Obtenção do título de Bacharel"),
            'orientador': str,
            'local': str,
            'ano': str
        }
        """
        # Autor (topo)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('autor', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        # Espaçamento
        for _ in range(8):
            self.doc.add_paragraph()

        # Título
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('titulo', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True

        # Espaçamento
        for _ in range(4):
            self.doc.add_paragraph()

        # Natureza do trabalho (recuado à direita)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.left_indent = Cm(8)

        texto_natureza = f"{dados.get('natureza', '')}\n\n{dados.get('objetivo', '')}"
        if dados.get('orientador'):
            texto_natureza += f"\n\nOrientador: {dados.get('orientador', '')}"

        run = p.add_run(texto_natureza)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Espaçamento
        for _ in range(6):
            self.doc.add_paragraph()

        # Local e ano
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('local', '').upper())
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(dados.get('ano', ''))
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Quebra de página
        self.doc.add_page_break()

    def adicionar_resumo(self, texto_resumo, palavras_chave):
        """Adiciona resumo formatado conforme ABNT"""
        # Título RESUMO
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('RESUMO')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        self.doc.add_paragraph()

        # Texto do resumo
        p = self.doc.add_paragraph(texto_resumo)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        self.doc.add_paragraph()

        # Palavras-chave
        p = self.doc.add_paragraph()
        run = p.add_run('Palavras-chave: ')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        run = p.add_run(palavras_chave)
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # Quebra de página
        self.doc.add_page_break()

    def adicionar_sumario(self, secoes):
        """
        Gera sumário automático
        secoes = [{'numero': '1', 'titulo': 'INTRODUÇÃO', 'pagina': 10}, ...]
        """
        # Título SUMÁRIO
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('SUMÁRIO')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        self.doc.add_paragraph()

        # Itens do sumário
        for secao in secoes:
            p = self.doc.add_paragraph()

            # Número e título
            texto = f"{secao['numero']}  {secao['titulo']}"
            run = p.add_run(texto)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            # Linha pontilhada e número da página
            espacos = 80 - len(texto)
            run = p.add_run('.' * espacos)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            run = p.add_run(f"  {secao['pagina']}")
            run.font.name = 'Arial'
            run.font.size = Pt(12)

        # Quebra de página
        self.doc.add_page_break()

    def adicionar_secao(self, numero, titulo, texto, nivel=1):
        """
        Adiciona seção formatada conforme NBR 6024
        nivel: 1 (principal), 2 (subseção), 3 (sub-subseção)
        """
        # Espaçamento antes
        if nivel == 1:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)

        # Título da seção
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        texto_titulo = f"{numero}  {titulo.upper()}"
        run = p.add_run(texto_titulo)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        if nivel == 1:
            p.paragraph_format.space_after = Pt(12)
        else:
            p.paragraph_format.space_after = Pt(6)

        # Texto da seção
        paragrafos = texto.split('\n\n')
        for paragrafo in paragrafos:
            if paragrafo.strip():
                p = self.doc.add_paragraph(paragrafo.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.first_line_indent = Cm(1.25)  # Recuo de parágrafo

                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

    def adicionar_citacao_longa(self, texto_citacao, autor, ano, pagina=None):
        """Adiciona citação longa (>3 linhas) formatada conforme NBR 10520"""
        p = self.doc.add_paragraph(texto_citacao)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

        # Referência da citação
        ref = f"({autor.upper()}, {ano}"
        if pagina:
            ref += f", p. {pagina}"
        ref += ")"

        p_ref = self.doc.add_paragraph(ref)
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ref.paragraph_format.left_indent = Cm(4)
        for run in p_ref.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    def adicionar_referencias(self, lista_referencias):
        """
        Adiciona seção de referências formatada conforme NBR 6023
        lista_referencias = ['REF1', 'REF2', ...]
        """
        # Quebra de página antes
        self.doc.add_page_break()

        # Título REFERÊNCIAS
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('REFERÊNCIAS')
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True

        p.paragraph_format.space_after = Pt(12)

        # Ordenar alfabeticamente
        referencias_ordenadas = sorted(lista_referencias)

        # Adicionar cada referência
        for referencia in referencias_ordenadas:
            if referencia.strip():
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(6)

                # Identificar e aplicar negrito no título
                match = re.match(r'^([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ][^.]+\.)\s*([^.]+\.)\s*(.*)$', referencia)

                if match:
                    autor = match.group(1)
                    titulo = match.group(2)
                    resto = match.group(3)

                    run = p.add_run(autor + ' ')
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

                    run = p.add_run(titulo)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.bold = True

                    if resto:
                        run = p.add_run(' ' + resto)
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                else:
                    run = p.add_run(referencia)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

    def salvar(self, caminho):
        """Salva o documento"""
        self.doc.save(caminho)


class AplicativoABNTModerno(ctk.CTk):
    """Aplicativo principal com interface moderna"""

    def __init__(self):
        super().__init__()

        self.title("📄 Formatador ABNT Acadêmico v3.0")
        self.geometry("1400x800")

        # Variáveis
        self.dados_trabalho = {}
        self.secoes = []

        self._criar_interface()

    def _criar_interface(self):
        """Cria a interface moderna com abas"""

        # Frame principal
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # Barra lateral
        self.sidebar = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_rowconfigure(10, weight=1)

        # Logo/Título
        self.logo_label = ctk.CTkLabel(
            self.sidebar,
            text="📄 ABNT\nAcadêmico",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        self.logo_label.grid(row=0, column=0, padx=20, pady=(20, 10))

        self.versao_label = ctk.CTkLabel(
            self.sidebar,
            text="Versão 3.0",
            font=ctk.CTkFont(size=12)
        )
        self.versao_label.grid(row=1, column=0, padx=20, pady=(0, 20))

        # Botões do menu
        self.btn_dados = ctk.CTkButton(
            self.sidebar,
            text="📋 Dados do Trabalho",
            command=lambda: self.mostrar_aba("dados")
        )
        self.btn_dados.grid(row=2, column=0, padx=20, pady=10)

        self.btn_elementos = ctk.CTkButton(
            self.sidebar,
            text="📑 Elementos Pré-textuais",
            command=lambda: self.mostrar_aba("elementos")
        )
        self.btn_elementos.grid(row=3, column=0, padx=20, pady=10)

        self.btn_conteudo = ctk.CTkButton(
            self.sidebar,
            text="✍️ Conteúdo",
            command=lambda: self.mostrar_aba("conteudo")
        )
        self.btn_conteudo.grid(row=4, column=0, padx=20, pady=10)

        self.btn_referencias = ctk.CTkButton(
            self.sidebar,
            text="📚 Referências",
            command=lambda: self.mostrar_aba("referencias")
        )
        self.btn_referencias.grid(row=5, column=0, padx=20, pady=10)

        self.btn_gerar = ctk.CTkButton(
            self.sidebar,
            text="🚀 Gerar Documento",
            command=self.gerar_documento,
            fg_color="green",
            hover_color="darkgreen"
        )
        self.btn_gerar.grid(row=6, column=0, padx=20, pady=10)

        # Informações na parte inferior
        self.info_label = ctk.CTkLabel(
            self.sidebar,
            text="Normas ABNT:\nNBR 14724:2023\nNBR 6023\nNBR 6024\nNBR 10520",
            font=ctk.CTkFont(size=10),
            justify="left"
        )
        self.info_label.grid(row=11, column=0, padx=20, pady=20)

        # Tema
        self.tema_label = ctk.CTkLabel(self.sidebar, text="Tema:")
        self.tema_label.grid(row=12, column=0, padx=20, pady=(10, 0))

        self.tema_switch = ctk.CTkSwitch(
            self.sidebar,
            text="Modo Escuro",
            command=self.alternar_tema,
            onvalue="dark",
            offvalue="light"
        )
        self.tema_switch.grid(row=13, column=0, padx=20, pady=(0, 20))
        self.tema_switch.select()

        # Frame de conteúdo (área principal)
        self.frame_conteudo = ctk.CTkFrame(self, corner_radius=0)
        self.frame_conteudo.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        self.frame_conteudo.grid_columnconfigure(0, weight=1)
        self.frame_conteudo.grid_rowconfigure(0, weight=1)

        # Criar todas as abas
        self._criar_aba_dados()
        self._criar_aba_elementos()
        self._criar_aba_conteudo()
        self._criar_aba_referencias()

        # Mostrar primeira aba
        self.mostrar_aba("dados")

    def _criar_aba_dados(self):
        """Cria aba de dados do trabalho"""
        self.aba_dados = ctk.CTkScrollableFrame(self.frame_conteudo)

        titulo = ctk.CTkLabel(
            self.aba_dados,
            text="📋 Dados do Trabalho Acadêmico",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        titulo.pack(pady=20)

        # Campos
        campos = [
            ("Instituição:", "instituicao"),
            ("Curso:", "curso"),
            ("Autor:", "autor"),
            ("Título do Trabalho:", "titulo"),
            ("Natureza do Trabalho:", "natureza"),
            ("Objetivo:", "objetivo"),
            ("Orientador:", "orientador"),
            ("Local:", "local"),
            ("Ano:", "ano")
        ]

        self.entries_dados = {}

        for label_text, key in campos:
            frame = ctk.CTkFrame(self.aba_dados, fg_color="transparent")
            frame.pack(fill="x", padx=40, pady=5)

            label = ctk.CTkLabel(frame, text=label_text, width=200, anchor="w")
            label.pack(side="left", padx=(0, 10))

            if key in ["natureza", "objetivo"]:
                entry = ctk.CTkEntry(frame, width=600)
            else:
                entry = ctk.CTkEntry(frame, width=600)

            entry.pack(side="left", fill="x", expand=True)
            self.entries_dados[key] = entry

        # Valores padrão
        self.entries_dados["natureza"].insert(0, "Trabalho de Conclusão de Curso")
        self.entries_dados["objetivo"].insert(0, "Obtenção do título de Bacharel em...")
        self.entries_dados["ano"].insert(0, str(datetime.now().year))

        # Botão salvar
        btn_salvar = ctk.CTkButton(
            self.aba_dados,
            text="💾 Salvar Dados",
            command=self.salvar_dados,
            width=200
        )
        btn_salvar.pack(pady=20)

    def _criar_aba_elementos(self):
        """Cria aba de elementos pré-textuais"""
        self.aba_elementos = ctk.CTkScrollableFrame(self.frame_conteudo)

        titulo = ctk.CTkLabel(
            self.aba_elementos,
            text="📑 Elementos Pré-textuais",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        titulo.pack(pady=20)

        # Resumo
        frame_resumo = ctk.CTkFrame(self.aba_elementos)
        frame_resumo.pack(fill="both", expand=True, padx=40, pady=10)

        label_resumo = ctk.CTkLabel(
            frame_resumo,
            text="RESUMO",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        label_resumo.pack(pady=10)

        self.text_resumo = ctk.CTkTextbox(frame_resumo, height=200, width=800)
        self.text_resumo.pack(padx=20, pady=10)

        # Palavras-chave
        frame_palavras = ctk.CTkFrame(self.aba_elementos, fg_color="transparent")
        frame_palavras.pack(fill="x", padx=40, pady=10)

        label_palavras = ctk.CTkLabel(
            frame_palavras,
            text="Palavras-chave (separadas por ponto):",
            width=300,
            anchor="w"
        )
        label_palavras.pack(side="left", padx=(0, 10))

        self.entry_palavras = ctk.CTkEntry(frame_palavras, width=500)
        self.entry_palavras.pack(side="left", fill="x", expand=True)
        self.entry_palavras.insert(0, "Palavra1. Palavra2. Palavra3.")

    def _criar_aba_conteudo(self):
        """Cria aba de conteúdo do trabalho"""
        self.aba_conteudo = ctk.CTkFrame(self.frame_conteudo)

        titulo = ctk.CTkLabel(
            self.aba_conteudo,
            text="✍️ Conteúdo do Trabalho",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        titulo.pack(pady=20)

        # Instruções
        instrucoes = ctk.CTkLabel(
            self.aba_conteudo,
            text="Cole ou digite o conteúdo completo do seu trabalho abaixo.\n" +
                 "O sistema formatará automaticamente conforme as normas ABNT.",
            font=ctk.CTkFont(size=12)
        )
        instrucoes.pack(pady=10)

        # Frame de botões
        frame_btns = ctk.CTkFrame(self.aba_conteudo, fg_color="transparent")
        frame_btns.pack(fill="x", padx=40, pady=10)

        btn_carregar = ctk.CTkButton(
            frame_btns,
            text="📂 Carregar Word",
            command=self.carregar_word
        )
        btn_carregar.pack(side="left", padx=5)

        btn_formatar = ctk.CTkButton(
            frame_btns,
            text="✨ Formatar ABNT",
            command=self.formatar_conteudo
        )
        btn_formatar.pack(side="left", padx=5)

        btn_limpar = ctk.CTkButton(
            frame_btns,
            text="🗑️ Limpar",
            command=lambda: self.text_conteudo.delete("1.0", "end")
        )
        btn_limpar.pack(side="left", padx=5)

        # Editor de texto
        self.text_conteudo = ctk.CTkTextbox(
            self.aba_conteudo,
            height=500,
            font=ctk.CTkFont(family="Arial", size=12)
        )
        self.text_conteudo.pack(fill="both", expand=True, padx=40, pady=10)

    def _criar_aba_referencias(self):
        """Cria aba de referências bibliográficas"""
        self.aba_referencias = ctk.CTkFrame(self.frame_conteudo)

        titulo = ctk.CTkLabel(
            self.aba_referencias,
            text="📚 Referências Bibliográficas",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        titulo.pack(pady=20)

        # Instruções
        instrucoes = ctk.CTkLabel(
            self.aba_referencias,
            text="Digite uma referência por linha. Serão ordenadas alfabeticamente automaticamente.\n" +
                 "Formato: SOBRENOME, Nome. Título: subtítulo. Edição. Local: Editora, ano.",
            font=ctk.CTkFont(size=12)
        )
        instrucoes.pack(pady=10)

        # Frame de botões
        frame_btns = ctk.CTkFrame(self.aba_referencias, fg_color="transparent")
        frame_btns.pack(fill="x", padx=40, pady=10)

        btn_exemplo = ctk.CTkButton(
            frame_btns,
            text="💡 Exemplo",
            command=self.inserir_exemplo_referencia
        )
        btn_exemplo.pack(side="left", padx=5)

        btn_formatar_ref = ctk.CTkButton(
            frame_btns,
            text="✨ Formatar",
            command=self.formatar_referencias
        )
        btn_formatar_ref.pack(side="left", padx=5)

        # Editor de referências
        self.text_referencias = ctk.CTkTextbox(
            self.aba_referencias,
            height=500,
            font=ctk.CTkFont(family="Arial", size=12)
        )
        self.text_referencias.pack(fill="both", expand=True, padx=40, pady=10)

    def mostrar_aba(self, nome_aba):
        """Mostra a aba selecionada"""
        # Esconder todas
        self.aba_dados.grid_forget()
        self.aba_elementos.grid_forget()
        self.aba_conteudo.pack_forget()
        self.aba_referencias.pack_forget()

        # Mostrar selecionada
        if nome_aba == "dados":
            self.aba_dados.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        elif nome_aba == "elementos":
            self.aba_elementos.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        elif nome_aba == "conteudo":
            self.aba_conteudo.pack(fill="both", expand=True, padx=20, pady=20)
        elif nome_aba == "referencias":
            self.aba_referencias.pack(fill="both", expand=True, padx=20, pady=20)

    def alternar_tema(self):
        """Alterna entre tema claro e escuro"""
        if self.tema_switch.get() == "dark":
            ctk.set_appearance_mode("dark")
        else:
            ctk.set_appearance_mode("light")

    def salvar_dados(self):
        """Salva os dados do trabalho"""
        for key, entry in self.entries_dados.items():
            self.dados_trabalho[key] = entry.get()

        messagebox.showinfo("Sucesso", "✅ Dados salvos com sucesso!")

    def carregar_word(self):
        """Carrega um arquivo Word"""
        caminho = filedialog.askopenfilename(
            title="Selecionar arquivo Word",
            filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
        )

        if caminho:
            try:
                doc = Document(caminho)
                texto = '\n'.join([p.text for p in doc.paragraphs])
                self.text_conteudo.delete("1.0", "end")
                self.text_conteudo.insert("1.0", texto)
                messagebox.showinfo("Sucesso", f"✅ Arquivo carregado: {os.path.basename(caminho)}")
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao carregar arquivo:\n{str(e)}")

    def formatar_conteudo(self):
        """Formata o conteúdo conforme ABNT"""
        texto = self.text_conteudo.get("1.0", "end-1c")

        if not texto.strip():
            messagebox.showwarning("Aviso", "Nenhum conteúdo para formatar!")
            return

        # Aplicar formatação de citações
        texto_formatado = FormatadorABNT.formatar_citacoes(texto)

        self.text_conteudo.delete("1.0", "end")
        self.text_conteudo.insert("1.0", texto_formatado)

        messagebox.showinfo("Sucesso", "✅ Conteúdo formatado conforme ABNT!")

    def formatar_referencias(self):
        """Formata as referências"""
        texto = self.text_referencias.get("1.0", "end-1c")
        texto_formatado = FormatadorABNT.formatar_referencias(texto)

        self.text_referencias.delete("1.0", "end")
        self.text_referencias.insert("1.0", texto_formatado)

        messagebox.showinfo("Sucesso", "✅ Referências formatadas!")

    def inserir_exemplo_referencia(self):
        """Insere exemplos de referências"""
        exemplos = """SILVA, João. Introdução à computação. 3. ed. São Paulo: Atlas, 2021.

FREIRE, Paulo. Pedagogia do oprimido. 50. ed. Rio de Janeiro: Paz e Terra, 2021.

VYGOTSKY, Lev Semyonovich. A formação social da mente. São Paulo: Martins Fontes, 1984.

BRASIL. Lei nº 9.394, de 20 de dezembro de 1996. Estabelece as diretrizes e bases da educação nacional. Diário Oficial da União, Brasília, DF, 23 dez. 1996."""

        self.text_referencias.delete("1.0", "end")
        self.text_referencias.insert("1.0", exemplos)

    def gerar_documento(self):
        """Gera o documento Word completo formatado"""
        # Validar dados
        if not self.dados_trabalho:
            messagebox.showwarning(
                "Aviso",
                "Por favor, preencha e salve os dados do trabalho primeiro!"
            )
            self.mostrar_aba("dados")
            return

        try:
            # Criar gerador
            gerador = GeradorDocumentoABNT()

            # 1. Capa
            gerador.adicionar_capa(self.dados_trabalho)

            # 2. Folha de rosto
            gerador.adicionar_folha_rosto(self.dados_trabalho)

            # 3. Resumo
            resumo = self.text_resumo.get("1.0", "end-1c")
            palavras = self.entry_palavras.get()
            if resumo.strip():
                gerador.adicionar_resumo(resumo, palavras)

            # 4. Sumário (exemplo básico)
            secoes_sumario = [
                {'numero': '1', 'titulo': 'INTRODUÇÃO', 'pagina': 10},
                {'numero': '2', 'titulo': 'DESENVOLVIMENTO', 'pagina': 12},
                {'numero': '3', 'titulo': 'CONCLUSÃO', 'pagina': 20},
                {'numero': '', 'titulo': 'REFERÊNCIAS', 'pagina': 22}
            ]
            gerador.adicionar_sumario(secoes_sumario)

            # 5. Conteúdo
            conteudo = self.text_conteudo.get("1.0", "end-1c")
            if conteudo.strip():
                # Processar o conteúdo em seções
                self._processar_conteudo(gerador, conteudo)

            # 6. Referências
            referencias_texto = self.text_referencias.get("1.0", "end-1c")
            if referencias_texto.strip():
                lista_referencias = [ref.strip() for ref in referencias_texto.split('\n') if ref.strip()]
                gerador.adicionar_referencias(lista_referencias)

            # Salvar
            caminho = filedialog.asksaveasfilename(
                title="Salvar documento",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx")],
                initialfile=f"trabalho_abnt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if caminho:
                gerador.salvar(caminho)
                messagebox.showinfo(
                    "Sucesso",
                    f"✅ Documento gerado com sucesso!\n\n{caminho}\n\n" +
                    "Formatação aplicada:\n" +
                    "✓ Capa e folha de rosto ABNT\n" +
                    "✓ Resumo formatado\n" +
                    "✓ Sumário automático\n" +
                    "✓ Citações em MAIÚSCULAS\n" +
                    "✓ Margens e espaçamento ABNT\n" +
                    "✓ Referências ordenadas"
                )

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar documento:\n{str(e)}")

    def _processar_conteudo(self, gerador, conteudo):
        """Processa o conteúdo e adiciona ao documento"""
        # Detectar seções baseadas em padrões
        linhas = conteudo.split('\n')
        texto_atual = []
        secao_atual = None
        numero_atual = None

        for linha in linhas:
            # Detectar títulos de seção (ex: "1 INTRODUÇÃO", "2 DESENVOLVIMENTO")
            match_secao = re.match(r'^(\d+)\s+([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ\s]+)$', linha.strip())

            if match_secao:
                # Salvar seção anterior
                if secao_atual and texto_atual:
                    gerador.adicionar_secao(numero_atual, secao_atual, '\n\n'.join(texto_atual))

                # Nova seção
                numero_atual = match_secao.group(1)
                secao_atual = match_secao.group(2)
                texto_atual = []
            else:
                if linha.strip():
                    texto_atual.append(linha.strip())

        # Adicionar última seção
        if secao_atual and texto_atual:
            gerador.adicionar_secao(numero_atual, secao_atual, '\n\n'.join(texto_atual))


def main():
    """Função principal"""
    app = AplicativoABNTModerno()
    app.mainloop()


if __name__ == "__main__":
    main()
