#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Formatador ABNT Desktop
Aplicativo para formatação automática de documentos conforme normas ABNT
Autor: Claude AI
Versão: 2.0 - Correções críticas aplicadas
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os
from datetime import datetime


class FormatadorCitacoes:
    """
    Classe responsável por formatar citações conforme ABNT NBR 10520:2023

    REGRA PRINCIPAL ABNT:
    - Citações entre parênteses: SOBRENOME em MAIÚSCULAS → (SILVA, 2022)
    - Citações fora de parênteses: Sobrenome com inicial maiúscula → Silva (2022)
    - et al.: sempre em minúsculas e itálico → et al.
    """

    @staticmethod
    def converter_citacoes_para_maiusculas(texto):
        """
        Converte citações entre parênteses para MAIÚSCULAS (padrão ABNT correto)

        Exemplos:
            (Silva, 2022) -> (SILVA, 2022)
            (santos, 2020) -> (SANTOS, 2020)
            (Junior, 2024) -> (JUNIOR, 2024)
            (Oliveira; Costa, 2021) -> (OLIVEIRA; COSTA, 2021)
        """

        def substituir_citacao_parenteses(match):
            conteudo_original = match.group(1)

            # Processa cada autor (separados por ponto e vírgula)
            autores = conteudo_original.split(';')
            autores_formatados = []

            for autor in autores:
                autor = autor.strip()

                # Separa o nome do ano/página
                partes = re.split(r'(,\s*\d{4})', autor, maxsplit=1)

                if len(partes) >= 2:
                    nome_parte = partes[0].strip()
                    resto = ''.join(partes[1:])  # ano e página se houver

                    # Processa cada palavra do nome
                    palavras = nome_parte.split()
                    palavras_maiusculas = []

                    for palavra in palavras:
                        # Mantém "et al." em minúsculas
                        if palavra.lower() in ['et', 'al', 'al.']:
                            palavras_maiusculas.append(palavra.lower())
                        else:
                            # Converte para MAIÚSCULAS
                            palavras_maiusculas.append(palavra.upper())

                    autor_formatado = ' '.join(palavras_maiusculas) + resto
                    autores_formatados.append(autor_formatado)
                else:
                    # Se não conseguiu separar, converte tudo para maiúsculas
                    palavras = autor.split()
                    palavras_maiusculas = []
                    for palavra in palavras:
                        if palavra.lower() in ['et', 'al', 'al.']:
                            palavras_maiusculas.append(palavra.lower())
                        else:
                            palavras_maiusculas.append(palavra.upper())
                    autores_formatados.append(' '.join(palavras_maiusculas))

            return f"({'; '.join(autores_formatados)})"

        # Padrão para citações entre parênteses com ano
        # Captura: (Nome, 2022) ou (Nome; Outro, 2022, p. 10)
        padrao = r'\(([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s,;]+\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)\)'
        texto = re.sub(padrao, substituir_citacao_parenteses, texto)

        return texto

    @staticmethod
    def converter_citacoes_fora_parenteses(texto):
        """
        Converte citações FORA de parênteses para maiúsculas também

        Exemplos:
            Segundo Silva (2022) -> Segundo SILVA (2022)
            Conforme Junior (2024, p. 15) -> Conforme JUNIOR (2024, p. 15)
        """

        def substituir_autor_fora(match):
            prefixo = match.group(1)  # "Segundo", "Conforme", etc.
            nome = match.group(2)
            citacao = match.group(3)  # (2022) ou (2024, p. 15)

            # Converte o nome para MAIÚSCULAS, exceto "et al."
            palavras = nome.split()
            palavras_maiusculas = []
            for palavra in palavras:
                if palavra.lower() in ['et', 'al', 'al.']:
                    palavras_maiusculas.append(palavra.lower())
                else:
                    palavras_maiusculas.append(palavra.upper())

            nome_formatado = ' '.join(palavras_maiusculas)

            return f"{prefixo} {nome_formatado} {citacao}"

        # Padrão: "Segundo/Conforme/De acordo com Nome (ano)"
        padrao = r'\b(Segundo|Conforme|De acordo com|Para)\s+([A-ZÀ-Ü][a-zà-ü]+(?:\s+[a-zà-ü]+)*)\s+(\(\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?\))'
        texto = re.sub(padrao, substituir_autor_fora, texto, flags=re.IGNORECASE)

        return texto

    @staticmethod
    def converter_multiplos_autores_para_et_al(texto):
        """
        Converte citações com 4+ autores para et al.
        Exemplo: (SANTOS; OLIVEIRA; COSTA; LIMA, 2020) -> (SANTOS et al., 2020)
        """

        def substituir(match):
            conteudo = match.group(1)

            # Conta quantos autores tem (separados por ponto e vírgula)
            autores = conteudo.split(';')

            # Se tem 4 ou mais autores, mantém apenas o primeiro + et al.
            if len(autores) >= 4:
                primeiro_autor = autores[0].strip()

                # Remove o ano do primeiro autor se existir
                primeiro_autor = re.sub(r',\s*\d{4}.*$', '', primeiro_autor).strip()

                # Pega o ano da citação original
                ano_match = re.search(r',\s*(\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)', conteudo)
                if ano_match:
                    ano = ano_match.group(1)
                    return f"({primeiro_autor} et al., {ano})"

            return match.group(0)

        # Aplica a conversão
        texto = re.sub(
            r'\(([A-ZÀ-Ü][A-Za-zÀ-ü]+(?:\s+[a-zà-ü]+)*(?:;\s*[A-ZÀ-Ü][A-Za-zÀ-ü]+(?:\s+[a-zà-ü]+)*){3,}[,\s]+\d{4}[a-z]?(?:,\s*p\.\s*\d+(?:-\d+)?)?)\)',
            substituir, texto
        )

        return texto

    @staticmethod
    def formatar_texto(texto):
        """Aplica todas as formatações de citações no texto"""

        # 1. Converte múltiplos autores (4+) para et al. PRIMEIRO
        texto = FormatadorCitacoes.converter_multiplos_autores_para_et_al(texto)

        # 2. Converte citações entre parênteses para MAIÚSCULAS
        texto = FormatadorCitacoes.converter_citacoes_para_maiusculas(texto)

        # 3. Converte citações fora de parênteses para MAIÚSCULAS
        texto = FormatadorCitacoes.converter_citacoes_fora_parenteses(texto)

        return texto


class FormatadorWord:
    """Classe responsável por aplicar formatação ABNT em documentos Word"""

    @staticmethod
    def criar_documento_formatado(texto):
        """
        Cria um documento Word com formatação ABNT completa:
        - Margens: 3cm (superior/esquerda), 2cm (inferior/direita)
        - Fonte: Arial 12
        - Espaçamento: 1,5
        - Alinhamento: Justificado
        - Resumo: Centralizado
        - Referências: Títulos em negrito
        """
        doc = Document()

        # Configurar margens (em Inches: 1 inch = 2.54 cm)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(3 / 2.54)      # 3 cm
            section.bottom_margin = Inches(2 / 2.54)    # 2 cm
            section.left_margin = Inches(3 / 2.54)      # 3 cm
            section.right_margin = Inches(2 / 2.54)     # 2 cm

        # Processar o texto em parágrafos
        paragrafos = texto.split('\n')

        i = 0
        while i < len(paragrafos):
            texto_paragrafo = paragrafos[i].strip()

            if not texto_paragrafo:
                # Parágrafo vazio (linha em branco)
                doc.add_paragraph()
                i += 1
                continue

            # Detectar tipo de parágrafo
            tipo = FormatadorWord._detectar_tipo_paragrafo(texto_paragrafo, paragrafos, i)

            if tipo == 'resumo_titulo':
                # Título "Resumo" centralizado
                FormatadorWord._formatar_titulo_resumo(doc, texto_paragrafo)

            elif tipo == 'titulo_secao':
                # Títulos principais (1 INTRODUÇÃO, 2 DESENVOLVIMENTO, etc.)
                FormatadorWord._formatar_titulo_secao(doc, texto_paragrafo)

            elif tipo == 'titulo_subsecao':
                # Subtítulos (2.1, 2.2, etc.) - CONVERTER PARA MAIÚSCULAS
                FormatadorWord._formatar_titulo_subsecao(doc, texto_paragrafo)

            elif tipo == 'citacao_longa':
                # Citação longa (>3 linhas)
                FormatadorWord._formatar_citacao_longa_paragrafo(doc, texto_paragrafo)

            elif tipo == 'referencia':
                # Referência bibliográfica (aplicar negrito no título)
                FormatadorWord._formatar_referencia(doc, texto_paragrafo)

            else:
                # Parágrafo normal
                FormatadorWord._formatar_paragrafo_padrao(doc, texto_paragrafo)

            i += 1

        return doc

    @staticmethod
    def _detectar_tipo_paragrafo(texto, todos_paragrafos, indice):
        """Detecta o tipo de parágrafo para aplicar formatação adequada"""

        texto_stripped = texto.strip()

        # Verifica se é o título "Resumo"
        if texto_stripped.lower() == 'resumo':
            return 'resumo_titulo'

        # Verifica se é título de seção (1 INTRODUÇÃO, 2 DESENVOLVIMENTO, REFERÊNCIAS)
        if re.match(r'^\d+\s+[A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ\s]+$', texto_stripped):
            return 'titulo_secao'

        if re.match(r'^[A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ\s]+$', texto_stripped) and len(texto_stripped.split()) <= 5:
            return 'titulo_secao'

        # Verifica se é subtítulo (2.1 Nome do Subtítulo)
        if re.match(r'^\d+\.\d+\s+.+', texto_stripped):
            return 'titulo_subsecao'

        # Verifica se é citação longa (>300 caracteres ou tem marcador específico)
        if len(texto_stripped) > 300 and ('"' in texto_stripped or '"' in texto_stripped):
            return 'citacao_longa'

        # Verifica se é referência (começa com SOBRENOME em maiúsculas seguido de ponto)
        if re.match(r'^[A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ][A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ\s,]+\.', texto_stripped):
            return 'referencia'

        return 'normal'

    @staticmethod
    def _formatar_titulo_resumo(doc, texto):
        """Formata o título 'Resumo' (centralizado, negrito, Arial 12)"""
        paragrafo = doc.add_paragraph(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in paragrafo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True

    @staticmethod
    def _formatar_titulo_secao(doc, texto):
        """Formata títulos de seção (1 INTRODUÇÃO, etc.) - negrito, maiúsculas"""
        paragrafo = doc.add_paragraph(texto.upper())
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragrafo_format = paragrafo.paragraph_format
        paragrafo_format.space_before = Pt(12)
        paragrafo_format.space_after = Pt(12)

        for run in paragrafo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True

    @staticmethod
    def _formatar_titulo_subsecao(doc, texto):
        """Formata subtítulos (2.1 Nome) - CONVERTER PARA MAIÚSCULAS"""
        # Extrai número e título
        match = re.match(r'^(\d+\.\d+)\s+(.+)$', texto)
        if match:
            numero = match.group(1)
            titulo = match.group(2)
            texto_formatado = f"{numero} {titulo.upper()}"
        else:
            texto_formatado = texto.upper()

        paragrafo = doc.add_paragraph(texto_formatado)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragrafo_format = paragrafo.paragraph_format
        paragrafo_format.space_before = Pt(12)
        paragrafo_format.space_after = Pt(6)

        for run in paragrafo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True

    @staticmethod
    def _formatar_citacao_longa_paragrafo(doc, texto):
        """Formata citação longa (recuo 4cm, fonte 10, espaçamento simples, SEM ASPAS)"""
        # Remove aspas do texto
        texto_limpo = texto.strip('"').strip('"').strip('"').strip()

        # IMPORTANTE: Preserva a citação se houver no final
        paragrafo = doc.add_paragraph(texto_limpo)

        # Alinhamento justificado
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Formato do parágrafo
        paragrafo_format = paragrafo.paragraph_format
        paragrafo_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Espaçamento simples
        paragrafo_format.left_indent = Inches(4 / 2.54)  # Recuo de 4 cm
        paragrafo_format.space_after = Pt(0)
        paragrafo_format.space_before = Pt(0)

        # Fonte Arial 10
        for run in paragrafo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

            # Formatar "et al." em itálico
            if 'et al.' in run.text:
                partes = run.text.split('et al.')
                run.text = partes[0]

                for j in range(1, len(partes)):
                    run_italic = paragrafo.add_run('et al.')
                    run_italic.italic = True
                    run_italic.font.name = 'Arial'
                    run_italic.font.size = Pt(10)

                    if partes[j]:
                        run_normal = paragrafo.add_run(partes[j])
                        run_normal.font.name = 'Arial'
                        run_normal.font.size = Pt(10)

    @staticmethod
    def _formatar_referencia(doc, texto):
        """
        Formata referência bibliográfica com título em NEGRITO

        Exemplo:
        GARTNER. Top Strategic Technology Trends 2023. ...
        O título "Top Strategic..." deve ficar em NEGRITO
        """
        paragrafo = doc.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragrafo_format = paragrafo.paragraph_format
        paragrafo_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragrafo_format.space_after = Pt(6)

        # Tenta identificar o título (geralmente entre o primeiro ponto e o segundo ponto, ou após autor e antes do local)
        # Padrão: AUTOR. Título da obra. Local: Editora, ano.

        # Regex para capturar: AUTOR. TÍTULO. Resto
        match = re.match(r'^([A-ZÀÁÂÃÄÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÖÙÚÛÜÝ][^.]+\.)\s*([^.]+\.)\s*(.*)$', texto)

        if match:
            autor = match.group(1)
            titulo = match.group(2)
            resto = match.group(3)

            # Adiciona autor (normal)
            run_autor = paragrafo.add_run(autor + ' ')
            run_autor.font.name = 'Arial'
            run_autor.font.size = Pt(12)

            # Adiciona título (NEGRITO)
            run_titulo = paragrafo.add_run(titulo)
            run_titulo.font.name = 'Arial'
            run_titulo.font.size = Pt(12)
            run_titulo.font.bold = True

            # Adiciona resto (normal)
            if resto:
                run_resto = paragrafo.add_run(' ' + resto)
                run_resto.font.name = 'Arial'
                run_resto.font.size = Pt(12)
        else:
            # Se não conseguir separar, formata tudo como texto normal
            run = paragrafo.add_run(texto)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

    @staticmethod
    def _formatar_paragrafo_padrao(doc, texto):
        """Aplica formatação ABNT padrão ao parágrafo"""
        paragrafo = doc.add_paragraph(texto)

        # Alinhamento justificado
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Formato do parágrafo
        paragrafo_format = paragrafo.paragraph_format
        paragrafo_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE  # 1,5
        paragrafo_format.space_after = Pt(0)
        paragrafo_format.space_before = Pt(0)

        # Fonte Arial 12
        for run in paragrafo.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            # Formatar "et al." em itálico
            if 'et al.' in run.text:
                partes = run.text.split('et al.')
                run.text = partes[0]

                for j in range(1, len(partes)):
                    run_italic = paragrafo.add_run('et al.')
                    run_italic.italic = True
                    run_italic.font.name = 'Arial'
                    run_italic.font.size = Pt(12)

                    if partes[j]:
                        run_normal = paragrafo.add_run(partes[j])
                        run_normal.font.name = 'Arial'
                        run_normal.font.size = Pt(12)

    @staticmethod
    def carregar_documento(caminho_arquivo):
        """Carrega um documento Word e retorna seu texto"""
        try:
            doc = Document(caminho_arquivo)
            texto_completo = []

            for paragrafo in doc.paragraphs:
                texto_completo.append(paragrafo.text)

            return '\n'.join(texto_completo)
        except Exception as e:
            raise Exception(f"Erro ao carregar documento: {str(e)}")

    @staticmethod
    def salvar_documento(doc, caminho_arquivo):
        """Salva o documento Word formatado"""
        try:
            doc.save(caminho_arquivo)
            return True
        except Exception as e:
            raise Exception(f"Erro ao salvar documento: {str(e)}")


class AplicativoFormatadorABNT:
    """Aplicativo principal com interface Tkinter"""

    def __init__(self, root):
        self.root = root
        self.root.title("📄 Formatador ABNT - Documentos (v2.0)")
        self.root.geometry("1200x700")

        # Variáveis
        self.texto_original = ""
        self.texto_formatado = ""

        # Configurar interface
        self._criar_interface()

    def _criar_interface(self):
        """Cria a interface gráfica completa"""

        # Frame superior - Botões de ação
        frame_botoes = ttk.Frame(self.root, padding="10")
        frame_botoes.pack(fill=tk.X)

        btn_carregar_word = ttk.Button(frame_botoes, text="📂 Carregar Word",
                                       command=self.carregar_word)
        btn_carregar_word.pack(side=tk.LEFT, padx=5)

        btn_inserir_texto = ttk.Button(frame_botoes, text="📝 Inserir Texto",
                                       command=self.inserir_texto)
        btn_inserir_texto.pack(side=tk.LEFT, padx=5)

        btn_formatar = ttk.Button(frame_botoes, text="✨ Formatar ABNT",
                                 command=self.formatar_texto,
                                 style='Accent.TButton')
        btn_formatar.pack(side=tk.LEFT, padx=5)

        btn_salvar = ttk.Button(frame_botoes, text="💾 Salvar Word",
                               command=self.salvar_word)
        btn_salvar.pack(side=tk.LEFT, padx=5)

        btn_limpar = ttk.Button(frame_botoes, text="🗑️ Limpar",
                               command=self.limpar_tudo)
        btn_limpar.pack(side=tk.LEFT, padx=5)

        # Separador
        ttk.Separator(self.root, orient=tk.HORIZONTAL).pack(fill=tk.X, pady=5)

        # Frame principal - Duas colunas
        frame_principal = ttk.Frame(self.root, padding="10")
        frame_principal.pack(fill=tk.BOTH, expand=True)

        # Coluna ANTES (Esquerda)
        frame_antes = ttk.LabelFrame(frame_principal, text="📄 ANTES (Original)",
                                     padding="10")
        frame_antes.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        self.texto_antes = scrolledtext.ScrolledText(frame_antes, wrap=tk.WORD,
                                                     font=("Arial", 10),
                                                     height=20)
        self.texto_antes.pack(fill=tk.BOTH, expand=True)

        btn_copiar_antes = ttk.Button(frame_antes, text="📋 Copiar Original",
                                     command=self.copiar_original)
        btn_copiar_antes.pack(pady=(5, 0))

        # Coluna DEPOIS (Direita)
        frame_depois = ttk.LabelFrame(frame_principal, text="✅ DEPOIS (Formatado ABNT)",
                                      padding="10")
        frame_depois.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        self.texto_depois = scrolledtext.ScrolledText(frame_depois, wrap=tk.WORD,
                                                      font=("Arial", 10),
                                                      height=20)
        self.texto_depois.pack(fill=tk.BOTH, expand=True)

        btn_copiar_depois = ttk.Button(frame_depois, text="📋 Copiar Formatado",
                                      command=self.copiar_formatado)
        btn_copiar_depois.pack(pady=(5, 0))

        # Barra de status
        self.status_bar = ttk.Label(self.root, text="✅ v2.0 - Correções aplicadas: citações em MAIÚSCULAS, resumo centralizado, títulos em negrito",
                                   relief=tk.SUNKEN, anchor=tk.W, padding="5")
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def carregar_word(self):
        """Carrega um arquivo Word (.docx)"""
        try:
            caminho = filedialog.askopenfilename(
                title="Selecionar arquivo Word",
                filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
            )

            if caminho:
                self.atualizar_status("Carregando documento...")
                texto = FormatadorWord.carregar_documento(caminho)
                self.texto_original = texto

                # Exibir no campo ANTES
                self.texto_antes.delete(1.0, tk.END)
                self.texto_antes.insert(1.0, texto)

                # Limpar campo DEPOIS
                self.texto_depois.delete(1.0, tk.END)

                self.atualizar_status(f"✅ Documento carregado: {os.path.basename(caminho)}")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar documento:\n{str(e)}")
            self.atualizar_status("❌ Erro ao carregar documento")

    def inserir_texto(self):
        """Permite inserir/editar texto diretamente no campo ANTES"""
        # Limpa o campo DEPOIS
        self.texto_depois.delete(1.0, tk.END)

        # Foca no campo ANTES para edição
        self.texto_antes.focus()

        self.atualizar_status("Digite ou cole o texto no campo 'ANTES' e clique em 'Formatar ABNT'")

    def formatar_texto(self):
        """Aplica formatação ABNT ao texto"""
        try:
            # Pega o texto do campo ANTES
            texto_original = self.texto_antes.get(1.0, tk.END).strip()

            if not texto_original:
                messagebox.showwarning("Aviso", "Nenhum texto para formatar!\n\nCarregue um documento ou insira texto primeiro.")
                return

            self.atualizar_status("Formatando texto conforme normas ABNT...")

            # Aplica formatação de citações
            texto_formatado = FormatadorCitacoes.formatar_texto(texto_original)

            self.texto_formatado = texto_formatado

            # Exibir no campo DEPOIS
            self.texto_depois.delete(1.0, tk.END)
            self.texto_depois.insert(1.0, texto_formatado)

            self.atualizar_status("✅ Documento formatado! Citações em MAIÚSCULAS conforme ABNT")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao formatar texto:\n{str(e)}")
            self.atualizar_status("❌ Erro na formatação")

    def copiar_original(self):
        """Copia o texto original para a área de transferência"""
        texto = self.texto_antes.get(1.0, tk.END).strip()
        if texto:
            self.root.clipboard_clear()
            self.root.clipboard_append(texto)
            self.atualizar_status("📋 Texto original copiado!")
        else:
            messagebox.showinfo("Aviso", "Nenhum texto para copiar")

    def copiar_formatado(self):
        """Copia o texto formatado para a área de transferência"""
        texto = self.texto_depois.get(1.0, tk.END).strip()
        if texto:
            self.root.clipboard_clear()
            self.root.clipboard_append(texto)
            self.atualizar_status("📋 Texto formatado copiado!")
        else:
            messagebox.showinfo("Aviso", "Nenhum texto formatado para copiar.\n\nClique em 'Formatar ABNT' primeiro.")

    def salvar_word(self):
        """Salva o documento formatado como .docx"""
        try:
            texto = self.texto_depois.get(1.0, tk.END).strip()

            if not texto:
                messagebox.showwarning("Aviso", "Nenhum texto formatado para salvar!\n\nClique em 'Formatar ABNT' primeiro.")
                return

            # Diálogo para escolher onde salvar
            caminho = filedialog.asksaveasfilename(
                title="Salvar documento formatado",
                defaultextension=".docx",
                filetypes=[("Documento Word", "*.docx"), ("Todos os arquivos", "*.*")],
                initialfile=f"documento_abnt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if caminho:
                self.atualizar_status("Salvando documento formatado...")

                # Criar documento com formatação ABNT completa
                doc = FormatadorWord.criar_documento_formatado(texto)

                # Salvar
                FormatadorWord.salvar_documento(doc, caminho)

                self.atualizar_status(f"✅ Documento salvo: {os.path.basename(caminho)}")
                messagebox.showinfo("Sucesso", f"Documento salvo com sucesso!\n\n{caminho}\n\nFormatação aplicada:\n✓ Citações em MAIÚSCULAS\n✓ Resumo centralizado\n✓ Títulos em negrito\n✓ Margens ABNT")

        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao salvar documento:\n{str(e)}")
            self.atualizar_status("❌ Erro ao salvar documento")

    def limpar_tudo(self):
        """Limpa todos os campos"""
        self.texto_antes.delete(1.0, tk.END)
        self.texto_depois.delete(1.0, tk.END)
        self.texto_original = ""
        self.texto_formatado = ""
        self.atualizar_status("Campos limpos. Pronto para novo documento.")

    def atualizar_status(self, mensagem):
        """Atualiza a barra de status"""
        self.status_bar.config(text=mensagem)
        self.root.update_idletasks()


def main():
    """Função principal para executar o aplicativo"""
    root = tk.Tk()
    app = AplicativoFormatadorABNT(root)
    root.mainloop()


if __name__ == "__main__":
    main()
